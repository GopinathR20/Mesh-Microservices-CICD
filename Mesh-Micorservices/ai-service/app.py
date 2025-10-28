import streamlit as st
import google.generativeai as genai
import os
import requests
import PyPDF2
import docx
from dotenv import load_dotenv
from io import BytesIO
import json
import time

# --- Configuration and Initialization ---
load_dotenv()
st.set_page_config(page_title="Mesh AI", page_icon="✨", layout="wide")

# --- Custom CSS for Styling ---
st.markdown("""
    <style>
    .main-title-container {
        display: flex;
        align-items: center;
        gap: 15px;
    }
    .main-title-container h1 {
        margin: 0;
        font-weight: 600;
    }
    .main-title-container span {
        font-size: 2.5rem;
    }
    .stButton>button {
        width: 100%;
        border-radius: 0.5rem;
    }
    .stExpander {
        border: 1px solid #333;
        border-radius: 0.5rem;
    }
    </style>
""", unsafe_allow_html=True)

# --- Gemini API Configuration ---
try:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        st.error("GEMINI_API_KEY environment variable not set.")
        model = None
    else:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
except Exception as e:
    st.error(f"Error configuring the AI model: {e}")
    model = None

# --- Helper Functions ---
@st.cache_data(show_spinner="Downloading and processing file...")
def get_text_from_url(url):
    """Downloads a file, extracts text, and caches it."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        file_bytes = BytesIO(response.content)
        text = ""
        if url.lower().endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file_bytes)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        elif url.lower().endswith('.docx'):
            document = docx.Document(file_bytes)
            for para in document.paragraphs:
                text += para.text + "\n"
        else:
            raise ValueError("Unsupported file type. Please use a .pdf or .docx URL.")
        
        if not text.strip():
            raise ValueError("Could not extract any text from the document.")
        return text
    except requests.exceptions.RequestException as e:
        raise ConnectionError(f"Could not download file. Please check the URL. Details: {e}")
    except Exception as e:
        raise RuntimeError(f"Failed to process the file. Details: {e}")

@st.cache_data(show_spinner="Processing uploaded file...")
def get_text_from_upload(uploaded_file):
    """Extracts text from an uploaded file object."""
    try:
        file_bytes = BytesIO(uploaded_file.getvalue())
        text = ""
        file_name = uploaded_file.name.lower()

        if file_name.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file_bytes)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        elif file_name.endswith('.docx'):
            document = docx.Document(file_bytes)
            for para in document.paragraphs:
                text += para.text + "\n"
        else:
            raise ValueError("Unsupported file type. Please upload a .pdf or .docx file.")
        
        if not text.strip():
            raise ValueError("Could not extract any text from the document.")
        return text
    except Exception as e:
        raise RuntimeError(f"Failed to process the uploaded file. Details: {e}")

def get_ai_response(prompt):
    """Generic function to get a response from the AI model."""
    if not model:
        return "AI model not available."
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"An error occurred with the AI model: {e}"

def generate_example_questions(document_text):
    """Generates example questions based on the document."""
    prompt = f"""
    Based on the following document text, generate exactly 3 concise and insightful questions that a user might ask.
    Return the questions as a JSON-formatted list of strings. For example: ["What is the main topic?", "Who are the key people mentioned?", "What is the conclusion?"]

    DOCUMENT TEXT:
    ---
    {document_text[:2000]}
    ---
    """
    response_text = get_ai_response(prompt)
    try:
        cleaned_response = response_text.strip().replace("```json", "").replace("```", "")
        questions = json.loads(cleaned_response)
        return questions if isinstance(questions, list) else []
    except (json.JSONDecodeError, TypeError):
        return []

def handle_question_click(question):
    """Callback to handle clicks on example question buttons."""
    st.session_state.user_question = question

def process_document(doc_text, source_id):
    """A centralized function to process text and populate session state."""
    st.session_state.clear()
    st.session_state.document_processed = True
    st.session_state.source_id = source_id
    st.session_state.document_text = doc_text
    st.session_state.summary = get_ai_response(f"Provide a concise summary of the following document. Use markdown for clear formatting:\n\n---\n\n{doc_text}")
    st.session_state.example_questions = generate_example_questions(doc_text)
    st.session_state.chat_history = []

# --- UI Rendering ---

# Header
st.markdown("""
    <div class="main-title-container">
        <h1>Mesh AI 🤖</h1>
    </div>
""", unsafe_allow_html=True)
st.write("Your personal AI assistant for asking questions about your documents.")

# --- Main Application Logic ---
if "document_processed" not in st.session_state:
    st.session_state.document_processed = False

file_url = st.query_params.get("file_url")

# If a URL is provided, process it automatically
if file_url and not st.session_state.get("source_id") == file_url:
    with st.spinner("Analyzing document and preparing your session... Please wait."):
        try:
            doc_text = get_text_from_url(file_url)
            process_document(doc_text, file_url)
        except (ValueError, ConnectionError, RuntimeError) as e:
            st.error(str(e))
            st.session_state.document_processed = False

# State 1: No document processed yet. Show input options.
if not st.session_state.document_processed:
    st.info("Analyze a document from a URL or by uploading a file.")
    tab1, tab2 = st.tabs(["🔗 Analyze from URL", "📤 Upload a Document"])

    with tab1:
        with st.form("url_form"):
            url_input = st.text_input("Enter public URL for a .pdf or .docx file", "")
            submitted = st.form_submit_button("Analyze Document")
            if submitted and url_input:
                st.query_params["file_url"] = url_input
                st.rerun()
            elif submitted:
                st.warning("Please enter a URL.")
    
    with tab2:
        uploaded_file = st.file_uploader("Choose a .pdf or .docx file", type=["pdf", "docx"])
        if uploaded_file is not None:
            with st.spinner("Analyzing document and preparing your session... Please wait."):
                try:
                    doc_text = get_text_from_upload(uploaded_file)
                    process_document(doc_text, uploaded_file.name)
                    st.rerun() 
                except (ValueError, RuntimeError) as e:
                    st.error(str(e))

# State 2: Document has been processed. Show the chat interface.
else:
    # Display Summary and Example Questions
    with st.expander("📄 Document Summary", expanded=True):
        st.markdown(st.session_state.get("summary", "No summary available."))
    
    st.markdown("##### Start the conversation")
    example_questions = st.session_state.get("example_questions", [])
    if example_questions:
        cols = st.columns(len(example_questions))
        for i, q in enumerate(example_questions):
            with cols[i]:
                st.button(q, on_click=handle_question_click, args=[q])
    st.markdown("---")

    # Display Chat History
    for message in st.session_state.get("chat_history", []):
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Handle user input
    user_question = st.chat_input("Or ask your own question...") or st.session_state.pop("user_question", None)

    if user_question:
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        with st.chat_message("user"):
            st.markdown(user_question)

        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                prompt = f"""
                Based ONLY on the document text provided below, answer the user's question. If the answer isn't in the document, state that clearly.
                DOCUMENT TEXT: --- {st.session_state.document_text} ---
                CONVERSATION HISTORY: --- {st.session_state.chat_history} ---
                USER'S QUESTION: {user_question}
                """
                response = get_ai_response(prompt)
                st.markdown(response)
        st.session_state.chat_history.append({"role": "assistant", "content": response})

    # Sidebar for actions
    with st.sidebar:
        st.header("Actions")
        if st.button("Analyze New Document"):
            st.query_params.clear()
            # Reset the session state to go back to the upload/URL screen
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
        if st.button("Clear Chat History"):
            st.session_state.chat_history = []
            st.rerun()

